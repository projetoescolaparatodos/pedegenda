from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

def configurar_abnt(doc):
    """Configura margens e estilo padrão ABNT"""
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    # Configurar fonte padrão para Arial (ou Times New Roman)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    # Espaçamento 1.5 e justificado é o padrão do corpo
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    return doc

def formatar_titulo(paragrafo, tamanho=12, negrito=True):
    run = paragrafo.runs[0]
    run.font.bold = negrito
    run.font.size = Pt(tamanho)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0, 0, 0) # Preto

def adicionar_texto_capa(doc, texto, negrito=True, tamanho=12, espaco_antes=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(espaco_antes)
    run = p.add_run(texto.upper())
    run.font.bold = negrito
    run.font.size = Pt(tamanho)
    run.font.name = 'Arial'

def criar_portfolio():
    doc = Document()
    doc = configurar_abnt(doc)

    # --- CAPA ---
    adicionar_texto_capa(doc, "UNIVERSIDADE FEDERAL DO PARÁ", tamanho=12)
    adicionar_texto_capa(doc, "CAMPUS UNIVERSITÁRIO DE ALTAMIRA", tamanho=12)
    adicionar_texto_capa(doc, "FACULDADE DE MEDICINA", tamanho=12)
    
    # Espaço para centralizar o nome
    adicionar_texto_capa(doc, "NALBERT HENRIQUE LUCAS RAMOS", tamanho=14, espaco_antes=150)
    
    # Espaço para o título
    adicionar_texto_capa(doc, "PORTFÓLIO ACADÊMICO: PRÁTICAS DE INTEGRAÇÃO ENSINO, SERVIÇO E COMUNIDADE (PIESC)", tamanho=14, espaco_antes=150)
    
    # Cidade e Ano no final
    adicionar_texto_capa(doc, "ALTAMIRA – PA", tamanho=12, espaco_antes=200)
    adicionar_texto_capa(doc, "2025", tamanho=12)
    
    doc.add_page_break()

    # --- FOLHA DE ROSTO ---
    adicionar_texto_capa(doc, "NALBERT HENRIQUE LUCAS RAMOS", tamanho=14)
    
    adicionar_texto_capa(doc, "PORTFÓLIO ACADÊMICO", tamanho=14, espaco_antes=150)
    
    # Texto de natureza do trabalho (Recuado)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(8) # Recuo da metade da página
    p.paragraph_format.line_spacing = 1.0  # Espaçamento simples
    p.paragraph_format.space_before = Pt(100)
    run = p.add_run("Portfólio acadêmico apresentado ao eixo Práticas de Integração Ensino, Serviço e Comunidade (PIESC) da Faculdade de Medicina da Universidade Federal do Pará (UFPA), como requisito avaliativo parcial do 2º semestre.")
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    
    # Preceptora
    p_orientador = doc.add_paragraph()
    p_orientador.paragraph_format.left_indent = Cm(8)
    p_orientador.paragraph_format.space_before = Pt(12)
    run_o = p_orientador.add_run("Preceptora: Enf. Yasmin Silva")
    run_o.font.name = 'Arial'
    run_o.font.size = Pt(10)

    # Cidade e Ano
    adicionar_texto_capa(doc, "ALTAMIRA – PA", tamanho=12, espaco_antes=150)
    adicionar_texto_capa(doc, "2025", tamanho=12)
    
    doc.add_page_break()

    # --- SUMÁRIO (Simulado) ---
    p = doc.add_paragraph("SUMÁRIO")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formatar_titulo(p, tamanho=12, negrito=True)
    
    itens_sumario = [
        "1. INTRODUÇÃO",
        "2. DIÁRIO DE ATIVIDADES",
        "3. APLICABILIDADE DO GENOGRAMA E ECOMAPA NA APS",
        "4. AÇÃO DE EDUCAÇÃO EM SAÚDE: DEZEMBRO VERMELHO",
        "5. AUTOAVALIAÇÃO E AVALIAÇÃO CONSTRUTIVA",
        "6. CONCLUSÃO",
        "REFERÊNCIAS",
        "APÊNDICES"
    ]
    
    for item in itens_sumario:
        p = doc.add_paragraph(item)
        p.paragraph_format.line_spacing = 1.5
    
    doc.add_page_break()

    # --- CONTEÚDO ---
    
    # Função auxiliar para texto corpo
    def add_body_text(documento, texto):
        p = documento.add_paragraph(texto)
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.line_spacing = 1.5
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return p

    # Função auxiliar para Títulos de Seção
    def add_heading(documento, texto):
        p = documento.add_paragraph(texto)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.runs[0]
        run.font.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        return p

    # 1. INTRODUÇÃO
    add_heading(doc, "1. INTRODUÇÃO")
    add_body_text(doc, "O portfólio acadêmico configura-se como uma metodologia ativa de ensino-aprendizagem que transcende o simples arquivamento de tarefas. Trata-se de um instrumento reflexivo que permite ao estudante de medicina documentar, analisar e reconstruir sua trajetória de formação. Conforme apontam Villas Boas (2004) e Sá-Chaves (2005), a construção do portfólio estimula a autonomia discente e o pensamento crítico, fundamentais para uma prática médica humanizada e contextualizada.")
    add_body_text(doc, "Este documento apresenta o registro descritivo e analítico das atividades desenvolvidas no eixo Práticas de Integração Ensino, Serviço e Comunidade (PIESC), realizadas durante o segundo semestre de 2025. As vivências ocorreram no território adscrito à Unidade Básica de Saúde (UBS) Santa Ana, em Altamira-PA, e incluíram visitas domiciliares, reconhecimento do território e ações de educação em saúde.")
    add_body_text(doc, "Além da descrição prática, este trabalho busca articular as vivências com o referencial teórico discutido nas conferências e aulas, abordando temas como abordagem familiar, indicadores de saúde, saúde mental e manejo de doenças crônicas. O objetivo central é demonstrar a compreensão da Atenção Primária à Saúde (APS) como ordenadora do cuidado e a importância da inserção precoce do estudante no Sistema Único de Saúde (SUS).")

    # 2. DIÁRIO
    add_heading(doc, "2. DIÁRIO DE ATIVIDADES")
    add_body_text(doc, "As atividades do semestre iniciaram-se com uma fundamentação teórica essencial para a prática no território. No dia 17 de outubro de 2025, foi realizada uma aula teórica sobre Abordagem Familiar. Este momento foi crucial para compreendermos que o adoecimento não ocorre isoladamente, mas é influenciado pela dinâmica e estrutura familiar. Segundo McGoldrick et al. (2012), ferramentas como o genograma são indispensáveis para visualizar padrões transgeracionais e o ciclo vital da família, preparando-nos para as visitas subsequentes.")
    add_body_text(doc, "Em 20 de outubro de 2025, iniciamos as práticas na UBS Santa Ana com uma visita domiciliar ao Sr. R.S. A atividade permitiu a aplicação prática da escuta qualificada. Durante a entrevista, observamos não apenas as queixas clínicas, mas o contexto social, identificando vínculos com a igreja e a vizinhança. Essa vivência reforçou os princípios da Política Nacional de Atenção Básica (PNAB), que preconiza o cuidado centrado na pessoa e situado em seu contexto sociocultural (BRASIL, 2012).")
    add_body_text(doc, "Dando continuidade à formação teórica, no dia 27 de outubro de 2025, participamos de uma conferência online abordando Educação em Saúde, Educação Popular em Saúde, Sistemas de Saúde e Indicadores da Atenção Básica. A discussão sobre indicadores e doenças de notificação compulsória foi particularmente relevante para entender como a gestão da UBS planeja suas ações baseadas em dados epidemiológicos locais, saindo do empirismo para uma gestão baseada em evidências.")
    add_body_text(doc, "No dia 03 de novembro de 2025, retornamos ao território para uma nova visita domiciliar. Desta vez, o foco foi a coleta detalhada de dados para a elaboração do genograma e ecomapa de uma família adscrita. A visita possibilitou compreender a complexidade das relações intrafamiliares e como a rede de apoio (ou a falta dela) impacta na adesão aos tratamentos propostos pela equipe de saúde.")
    add_body_text(doc, "Nas semanas de 10 e 17 de novembro de 2025, as atividades práticas no território foram suspensas devido à realização da COP30 na região, o que inviabilizou o deslocamento e a logística das equipes. Contudo, esse tempo foi utilizado para a sistematização dos dados coletados anteriormente.")
    add_body_text(doc, "Retomamos as atividades presenciais em 24 de novembro de 2025, com a entrega e discussão dos genogramas e ecomapas elaborados. Apresentamos os instrumentos à Agente Comunitária de Saúde (ACS) responsável pela microárea. Essa devolutiva é fundamental para a equipe, pois, conforme Wright e Leahey (2012), esses instrumentos visuais facilitam a identificação rápida de vulnerabilidades e potencialidades da família, auxiliando a equipe da ESF no planejamento do cuidado longitudinal.")
    add_body_text(doc, "No dia 12 de dezembro de 2025, participamos de uma conferência teórica sobre a Política Nacional de Atenção às Urgências e Emergências (PNAU). Embora nosso foco seja a APS, compreender a rede de urgência é vital para entender os fluxos de referência e contrarreferência dentro da Rede de Atenção à Saúde (RAS).")
    add_body_text(doc, "Em 15 de dezembro de 2025, executamos a ação de intervenção comunitária na UBS Santa Ana, alusiva ao \"Dezembro Vermelho\". A atividade será detalhada em seção específica deste portfólio, mas destaca-se a relevância de trabalhar a prevenção combinada do HIV/ISTs no território.")
    add_body_text(doc, "Finalizando o ciclo de aprendizado teórico, tivemos duas aulas online importantes. Em 16 de dezembro de 2025, o tema foi Políticas e Ações de Controle do Diabetes Mellitus, essencial dada a alta prevalência desta condição na comunidade visitada. Já em 17 de dezembro de 2025, discutiu-se a Política Nacional de Saúde Mental, onde compreendemos a articulação entre a Atenção Básica e os Centros de Atenção Psicossocial (CAPS) no manejo do sofrimento psíquico, pautado na reforma psiquiátrica e na luta antimanicomial.")

    # 3. GENOGRAMA
    add_heading(doc, "3. APLICABILIDADE DO GENOGRAMA E ECOMAPA NA APS")
    add_body_text(doc, "Durante o semestre, a construção do genograma e do ecomapa consolidou-se como o eixo central da avaliação familiar. O genograma permitiu a representação gráfica de três gerações da família acompanhada, evidenciando a recorrência de doenças crônicas (como hipertensão e diabetes) e a natureza dos vínculos afetivos (conflituosos ou harmoniosos). Já o ecomapa ilustrou as conexões da família com o meio externo, revelando a importância de instituições como a igreja e a própria UBS como fontes de apoio social.")
    add_body_text(doc, "As representações gráficas elaboradas encontram-se nos apêndices deste trabalho. A construção desses diagramas confirmou que a abordagem familiar é indispensável para um plano terapêutico eficaz, pois permite identificar quem são os cuidadores principais e quais recursos comunitários podem ser mobilizados em prol do paciente.")

    # 4. EDUCAÇÃO EM SAÚDE
    add_heading(doc, "4. AÇÃO DE EDUCAÇÃO EM SAÚDE: DEZEMBRO VERMELHO")
    add_body_text(doc, "A atividade de educação em saúde foi realizada na UBS Santa Ana no dia 15 de dezembro de 2025, sob supervisão da preceptora Yasmin Silva. O tema escolhido foi \"Dezembro Vermelho: conscientização, prevenção e direitos humanos sobre HIV/AIDS e ISTs\".")
    
    p = doc.add_paragraph("Planejamento e Metodologia")
    p.runs[0].font.bold = True
    p.paragraph_format.first_line_indent = Cm(1.25)
    
    add_body_text(doc, "Identificamos como problemas principais no território o diagnóstico tardio do HIV e o estigma social que afasta a população da testagem. Para abordar o tema de forma leve e acessível a diferentes faixas etárias, optamos pela gamificação através da dinâmica do \"Baralho Mito ou Verdade\".")
    add_body_text(doc, "Os acadêmicos confeccionaram cartas contendo afirmações populares sobre transmissão e prevenção. Os participantes retiravam uma carta, opinavam se era mito ou verdade, e em seguida a equipe esclarecia a questão com embasamento científico, abordando a diferença entre HIV e AIDS, a importância da PEP/PrEP e o funcionamento dos testes rápidos disponíveis na unidade.")

    p = doc.add_paragraph("Resultados e Reflexão")
    p.runs[0].font.bold = True
    p.paragraph_format.first_line_indent = Cm(1.25)
    
    add_body_text(doc, "A metodologia ativa favoreceu a interação, permitindo que dúvidas fossem sanadas sem constrangimentos. Percebeu-se que o formato lúdico foi eficaz para desmistificar preconceitos arraigados. A ação reforçou o papel do médico não apenas como clínico, mas como educador em saúde, capaz de traduzir o conhecimento técnico para uma linguagem acessível, empoderando a comunidade para o autocuidado.")

    # 5. AUTOAVALIAÇÃO
    add_heading(doc, "5. AUTOAVALIAÇÃO E AVALIAÇÃO CONSTRUTIVA")
    add_body_text(doc, "O percurso no eixo PIESC foi desafiador e enriquecedor. Considero meu desempenho satisfatório, tendo participado ativamente de todas as visitas, discussões teóricas e planejamentos. A principal competência desenvolvida foi o olhar ampliado: aprendi a enxergar o paciente para além da patologia, considerando seu contexto familiar e social.")
    add_body_text(doc, "O trabalho em equipe fluiu com colaboração mútua, superando as dificuldades logísticas impostas pelo calendário e eventos externos (COP30). A preceptoria da Enfermeira Yasmin foi essencial para guiar nossas intervenções e garantir a segurança nas práticas. Como ponto de melhoria, identifico a necessidade de aprofundar o conhecimento sobre os fluxos burocráticos da rede de saúde para agilizar encaminhamentos futuros.")

    # 6. CONCLUSÃO
    add_heading(doc, "6. CONCLUSÃO")
    add_body_text(doc, "A elaboração deste portfólio permitiu a síntese reflexiva das experiências vivenciadas, conectando a teoria da sala de aula com a realidade da UBS Santa Ana. As atividades evidenciaram que a Atenção Primária é complexa e exige do médico competências que vão além da técnica, incluindo comunicação, empatia e conhecimento de gestão em saúde.")
    add_body_text(doc, "O uso de ferramentas como o genograma e a realização de ações educativas demonstraram, na prática, os princípios do SUS: integralidade, equidade e participação popular. Conclui-se que o semestre foi fundamental para a construção de uma identidade profissional comprometida com a saúde pública e com a realidade social da população amazônica.")

    # REFERÊNCIAS
    doc.add_page_break()
    p = doc.add_paragraph("REFERÊNCIAS")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT # ABNT pede alinhado à esquerda ou centralizado, geralmente esquerda na lista
    formatar_titulo(p, 12, True)

    referencias = [
        "BRASIL. Ministério da Saúde. Política Nacional de Atenção Básica. Brasília: Ministério da Saúde, 2012.",
        "BRASIL. Ministério da Saúde. Diretrizes Curriculares Nacionais do Curso de Graduação em Medicina. Brasília: MEC, 2014.",
        "BRASIL. Ministério da Saúde. Protocolo clínico e diretrizes terapêuticas para manejo da infecção pelo HIV em adultos. Brasília: Ministério da Saúde, 2023. Disponível em: <https://www.gov.br/saude/pt-br/assuntos/saude-de-a-a-z/h/hiv-aids>. Acesso em: 15 dez. 2025.",
        "BRASIL. Ministério da Saúde. Infecções sexualmente transmissíveis. Brasília: Ministério da Saúde, 2023. Disponível em: <https://www.gov.br/saude/pt-br/assuntos/saude-de-a-a-z/i/ist>. Acesso em: 15 dez. 2025.",
        "MCGOLDRICK, M.; GERSON, R.; PETRY, S. Genogramas: avaliação e intervenção familiar. 3. ed. Porto Alegre: Artmed, 2012.",
        "SÁ-CHAVES, I. Portfólios reflexivos: estratégia de formação e de supervisão. Porto: Porto Editora, 2005.",
        "VILLAS BOAS, B. M. F. Portfólio, avaliação e trabalho pedagógico. Campinas: Papirus, 2004.",
        "WRIGHT, L. M.; LEAHEY, M. Enfermeiras e famílias: um guia para avaliação e intervenção na família. 5. ed. São Paulo: Roca, 2012."
    ]

    for ref in referencias:
        p = doc.add_paragraph(ref)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # APÊNDICES
    doc.add_page_break()
    p = doc.add_paragraph("APÊNDICES")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formatar_titulo(p, 12, True)

    def add_appendix_placeholder(doc, titulo, legenda):
        p = doc.add_paragraph(titulo)
        p.runs[0].font.bold = True
        p.paragraph_format.space_before = Pt(24)
        
        # Placeholder da imagem
        p_img = doc.add_paragraph("[INSERIR IMAGEM AQUI]")
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Legenda
        p_leg = doc.add_paragraph(legenda)
        p_leg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_leg.runs[0].font.size = Pt(10)

    add_appendix_placeholder(doc, 
                             "APÊNDICE A – REGISTRO FOTOGRÁFICO DAS VISITAS DOMICILIARES", 
                             "Fonte: Acervo pessoal (2025).")

    add_appendix_placeholder(doc, 
                             "APÊNDICE B – GENOGRAMAS E ECOMAPAS DOS PACIENTES ACOMPANHADOS", 
                             "Fonte: Elaborado pelos autores (2025).")

    add_appendix_placeholder(doc, 
                             "APÊNDICE C – REGISTROS DA AÇÃO DE EDUCAÇÃO EM SAÚDE", 
                             "Fonte: Acervo pessoal (2025).")

    # Salvar
    doc.save("Portfolio_Nalbert_ABNT.docx")
    print("Arquivo 'Portfolio_Nalbert_ABNT.docx' criado com sucesso!")

if __name__ == "__main__":
    criar_portfolio()
